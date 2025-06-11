"""
Template system for generating realistic chaff documents.

This module provides classes for generating various file types with realistic
content including PDFs, Word documents, Excel files, emails, and more.
"""

import os
import random
import csv
import io
import math
import tempfile
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
from pathlib import Path

# Third-party imports
from faker import Faker
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from PIL import Image, ImageDraw, ImageFont

# Import phrase generator
try:
    from phrase_generator import PhraseGenerator
    PHRASE_GENERATOR_AVAILABLE = True
except ImportError:
    PHRASE_GENERATOR_AVAILABLE = False


class BaseTemplate:
    """Base class for all document templates"""
    
    def __init__(self, language: str = 'en'):
        self.language = language
        # Map unsupported language codes to supported ones
        faker_locale_map = {
            'cn': 'zh_CN',  # Chinese
            'jp': 'ja_JP',  # Japanese
            'en': 'en_US',  # English
            'es': 'es_ES',  # Spanish
            'fr': 'fr_FR',  # French
            'de': 'de_DE',  # German
            'ru': 'ru_RU'   # Russian
        }
        faker_locale = faker_locale_map.get(language, 'en_US')
        self.fake = Faker(faker_locale)
        
        # Initialize phrase generator if available
        if PHRASE_GENERATOR_AVAILABLE:
            self.phrase_gen = PhraseGenerator()
        else:
            self.phrase_gen = None
        
        # Set up randomized date ranges for realistic document timestamps
        self._setup_random_date_ranges()
    
    def _setup_random_date_ranges(self):
        """Setup randomized date ranges for document content"""
        from datetime import datetime, timedelta
        import random
        
        now = datetime.now()
        
        # Create realistic date ranges for different document ages
        age_ranges = {
            'recent': (now - timedelta(days=90), now - timedelta(days=1)),
            'medium': (now - timedelta(days=730), now - timedelta(days=90)),  # 3 months to 2 years
            'old': (now - timedelta(days=1825), now - timedelta(days=730)),   # 2-5 years
            'archive': (now - timedelta(days=3650), now - timedelta(days=1825))  # 5-10 years
        }
        
        # Randomly select an age category for this document
        age_category = random.choices(
            list(age_ranges.keys()),
            weights=[0.3, 0.4, 0.2, 0.1]  # More recent documents are more common
        )[0]
        
        self.doc_date_range = age_ranges[age_category]
        
        # Generate a base document date within this range
        start_date, end_date = self.doc_date_range
        time_span = (end_date - start_date).total_seconds()
        random_offset = random.uniform(0, time_span)
        self.base_document_date = start_date + timedelta(seconds=random_offset)
    
    def get_random_document_date(self, offset_days: int = 0) -> datetime:
        """Get a randomized date for document content, optionally offset from base date"""
        from datetime import timedelta
        return self.base_document_date + timedelta(days=offset_days)
    
    def get_random_date_in_range(self, days_before: int = 365, days_after: int = 30) -> datetime:
        """Get a random date within a range relative to the base document date"""
        from datetime import timedelta
        import random
        
        start_date = self.base_document_date - timedelta(days=days_before)
        end_date = self.base_document_date + timedelta(days=days_after)
        
        time_span = (end_date - start_date).total_seconds()
        random_offset = random.uniform(0, time_span)
        return start_date + timedelta(seconds=random_offset)
        
    def generate(self, target_size: int, filename: str) -> bytes:
        """Generate document content. Must be implemented by subclasses."""
        raise NotImplementedError("Subclasses must implement generate method")
    
    def get_realistic_content(self, content_type: str) -> str:
        """Generate realistic content based on type"""
        if content_type == 'business_letter':
            return self._generate_business_letter()
        elif content_type == 'report':
            return self._generate_report()
        elif content_type == 'invoice':
            return self._generate_invoice()
        elif content_type == 'memo':
            return self._generate_memo()
        else:
            return self._generate_generic_text()
    
    def _generate_business_letter(self) -> str:
        company = self.fake.company()
        name = self.fake.name()
        address = self.fake.address().replace('\n', ', ')
        letter_date = self.get_random_document_date()
        
        return f"""
{company}
{address}

{letter_date.strftime('%B %d, %Y')}

Dear {name},

{self.fake.text(max_nb_chars=500)}

We appreciate your continued business and look forward to working with you.

Sincerely,

{self.fake.name()}
{self.fake.job()}
{company}
"""
    
    def _generate_report(self) -> str:
        report_date = self.get_random_document_date()
        return f"""
QUARTERLY BUSINESS REPORT
{report_date.strftime('%B %d, %Y')}

EXECUTIVE SUMMARY
{self.fake.text(max_nb_chars=300)}

KEY METRICS
- Revenue: ${self.fake.random_int(min=100000, max=5000000):,}
- Growth: {self.fake.random_int(min=-10, max=50)}%
- Customer Satisfaction: {self.fake.random_int(min=70, max=98)}%

ANALYSIS
{self.fake.text(max_nb_chars=800)}

RECOMMENDATIONS
{self.fake.text(max_nb_chars=400)}
"""
    
    def _generate_invoice(self) -> str:
        invoice_date = self.get_random_document_date()
        due_date = self.get_random_date_in_range(days_before=0, days_after=30)
        
        return f"""
INVOICE #{self.fake.random_int(min=1000, max=9999)}

Bill To:
{self.fake.name()}
{self.fake.company()}
{self.fake.address().replace(chr(10), ', ')}

Date: {invoice_date.strftime('%B %d, %Y')}
Due Date: {due_date.strftime('%B %d, %Y')}

Items:
- {self.fake.catch_phrase()}: ${self.fake.random_int(min=100, max=5000)}
- {self.fake.bs()}: ${self.fake.random_int(min=50, max=2000)}

Subtotal: ${self.fake.random_int(min=500, max=10000)}
Tax: ${self.fake.random_int(min=50, max=1000)}
Total: ${self.fake.random_int(min=600, max=11000)}
"""
    
    def _generate_memo(self) -> str:
        memo_date = self.get_random_document_date()
        
        return f"""
MEMORANDUM

TO: All Staff
FROM: {self.fake.name()}, {self.fake.job()}
DATE: {memo_date.strftime('%B %d, %Y')}
RE: {self.fake.catch_phrase()}

{self.fake.text(max_nb_chars=600)}

Please contact me if you have any questions.

{self.fake.name()}
"""
    
    def _generate_generic_text(self) -> str:
        return self.fake.text(max_nb_chars=1000)


class PDFTemplate(BaseTemplate):
    """Generate realistic PDF documents"""
    
    def __init__(self, language: str = 'en'):
        super().__init__(language)
        self.image_refs = []  # Store image references for embedding
    
    def set_image_references(self, image_refs: List[str]):
        """Set list of image filenames to embed in this document"""
        self.image_refs = image_refs
    
    def generate(self, target_size: int, filename: str) -> bytes:
        buffer = io.BytesIO()
        
        # Create PDF
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Choose document type
        doc_types = ['business_letter', 'report', 'invoice', 'memo']
        doc_type = random.choice(doc_types)
        content = self.get_realistic_content(doc_type)
        
        # Add title
        c.setFont("Helvetica-Bold", 16)
        title = f"{doc_type.replace('_', ' ').title()}"
        c.drawString(50, height - 50, title)
        
        # Add content
        c.setFont("Helvetica", 12)
        y_position = height - 100
        
        content_lines = content.split('\n')
        
        for i, line in enumerate(content_lines):
            if y_position < 50:  # Start new page
                c.showPage()
                y_position = height - 50
                c.setFont("Helvetica", 12)
            
            # Add image references randomly throughout content
            if self.image_refs and random.random() < 0.3 and y_position > 100:
                image_ref = random.choice(self.image_refs)
                
                # Draw image placeholder box
                c.setStrokeColorRGB(0.5, 0.5, 0.5)
                c.rect(50, y_position - 60, 200, 50, stroke=1, fill=0)
                
                # Add image reference text
                c.setFont("Helvetica-Bold", 10)
                c.drawString(55, y_position - 30, f"[IMAGE: {image_ref}]")
                c.setFont("Helvetica", 9)
                c.drawString(55, y_position - 45, f"Visual content: {image_ref}")
                
                y_position -= 80
                c.setFont("Helvetica", 12)
            
            # Wrap long lines
            if len(line) > 80:
                words = line.split(' ')
                current_line = ''
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + ' '
                    else:
                        c.drawString(50, y_position, current_line.strip())
                        y_position -= 15
                        current_line = word + ' '
                        if y_position < 50:
                            c.showPage()
                            y_position = height - 50
                if current_line:
                    c.drawString(50, y_position, current_line.strip())
                    y_position -= 15
            else:
                c.drawString(50, y_position, line)
                y_position -= 15
        
        # Add image references section if we have them
        if self.image_refs and y_position > 150:
            y_position -= 20
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y_position, "Referenced Images:")
            y_position -= 20
            
            c.setFont("Helvetica", 10)
            for image_ref in self.image_refs:
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                c.drawString(50, y_position, f"• {image_ref} - Embedded above")
                y_position -= 12
        
        # Add footer with randomized date
        c.setFont("Helvetica", 8)
        footer_date = self.get_random_document_date()
        c.drawString(50, 30, f"Generated on {footer_date.strftime('%Y-%m-%d %H:%M')}")
        
        c.save()
        content_bytes = buffer.getvalue()
        
        # For PDFs, don't pad with binary data as it corrupts the format
        # If we need more content, generate a new PDF with more pages
        if len(content_bytes) < target_size * 0.8 and target_size > 50000:
            # Create a new PDF with more content
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            # Add multiple pages of content
            pages_needed = max(2, (target_size // 50000))
            for page_num in range(min(pages_needed, 10)):  # Cap at 10 pages
                c.setFont("Helvetica-Bold", 16)
                c.drawString(50, height - 50, f"{doc_type.replace('_', ' ').title()} - Page {page_num + 1}")
                
                c.setFont("Helvetica", 12)
                y_position = height - 100
                
                # Add more content for this page
                if self.phrase_gen:
                    page_content = self.phrase_gen.generate_realistic_document_content(2000)
                else:
                    page_content = self.get_realistic_content(doc_type)
                
                for line in page_content.split('\n'):
                    if y_position < 50:
                        break
                    
                    # Add occasional image references in additional pages
                    if self.image_refs and random.random() < 0.2 and y_position > 100:
                        image_ref = random.choice(self.image_refs)
                        c.setStrokeColorRGB(0.5, 0.5, 0.5)
                        c.rect(50, y_position - 40, 150, 30, stroke=1, fill=0)
                        c.setFont("Helvetica-Bold", 9)
                        c.drawString(55, y_position - 25, f"[REF: {image_ref}]")
                        y_position -= 50
                        c.setFont("Helvetica", 12)
                    
                    if len(line) > 80:
                        words = line.split(' ')
                        current_line = ''
                        for word in words:
                            if len(current_line + word) < 80:
                                current_line += word + ' '
                            else:
                                c.drawString(50, y_position, current_line.strip())
                                y_position -= 15
                                current_line = word + ' '
                                if y_position < 50:
                                    break
                        if current_line and y_position >= 50:
                            c.drawString(50, y_position, current_line.strip())
                            y_position -= 15
                    else:
                        c.drawString(50, y_position, line)
                        y_position -= 15
                
                c.setFont("Helvetica", 8)
                footer_date = self.get_random_document_date()
                c.drawString(50, 30, f"Generated on {footer_date.strftime('%Y-%m-%d %H:%M')} - Page {page_num + 1}")
                
                if page_num < pages_needed - 1:
                    c.showPage()
            
            c.save()
            content_bytes = buffer.getvalue()
        
        return content_bytes


class DOCXTemplate(BaseTemplate):
    """Generate realistic Word documents"""
    
    def __init__(self, language: str = 'en'):
        super().__init__(language)
        self.image_refs = []  # Store image references for embedding
    
    def set_image_references(self, image_refs: List[str]):
        """Set list of image filenames to embed in this document"""
        self.image_refs = image_refs
    
    def generate(self, target_size: int, filename: str) -> bytes:
        doc = Document()
        
        # Choose document type
        doc_types = ['business_letter', 'report', 'invoice', 'memo']
        doc_type = random.choice(doc_types)
        
        # Add title
        title = doc.add_heading(doc_type.replace('_', ' ').title(), 0)
        
        # Add content
        content = self.get_realistic_content(doc_type)
        
        paragraphs = content.split('\n\n')
        for i, paragraph_text in enumerate(paragraphs):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
                
                # Randomly embed image placeholders after some paragraphs
                if self.image_refs and random.random() < 0.4:  # 40% chance
                    image_ref = random.choice(self.image_refs)
                    
                    # Add image placeholder paragraph
                    image_para = doc.add_paragraph()
                    image_para.add_run(f"[IMAGE: {image_ref}]").bold = True
                    
                    # Add image description
                    description_para = doc.add_paragraph()
                    descriptions = [
                        f"Figure: {image_ref} - Supporting visual documentation",
                        f"Image Reference: {image_ref} contains relevant charts and diagrams",
                        f"Visual Aid: See {image_ref} for detailed illustrations",
                        f"Embedded Content: {image_ref} provides additional context",
                        f"Attachment: {image_ref} - graphical representation of data"
                    ]
                    description_para.add_run(random.choice(descriptions)).italic = True
        
        # Add a table for reports and invoices
        if doc_type in ['report', 'invoice']:
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Item'
            hdr_cells[1].text = 'Quantity'
            hdr_cells[2].text = 'Amount'
            
            for _ in range(random.randint(3, 8)):
                row_cells = table.add_row().cells
                row_cells[0].text = self.fake.word()
                row_cells[1].text = str(self.fake.random_int(min=1, max=100))
                row_cells[2].text = f"${self.fake.random_int(min=10, max=1000)}"
        
        # Add section for embedded images if we have them
        if self.image_refs:
            doc.add_heading('Visual References', level=2)
            for image_ref in self.image_refs:
                doc.add_paragraph(f"• {image_ref} - Embedded in document above")
        
        # Add more content if we need to reach target size
        if target_size > 100000:  # Only for large target sizes
            # Add more paragraphs to reach approximate target size
            current_size = 0
            attempts = 0
            while current_size < target_size * 0.8 and attempts < 50:  # Don't try forever
                if self.phrase_gen:
                    additional_content = self.phrase_gen.generate_realistic_document_content(1000)
                else:
                    additional_content = self.fake.text(max_nb_chars=1000)
                
                doc.add_paragraph(additional_content)
                
                # Occasionally add more image references in additional content
                if self.image_refs and random.random() < 0.2:
                    image_ref = random.choice(self.image_refs)
                    doc.add_paragraph(f"[REFERENCE: {image_ref}]").runs[0].bold = True
                
                # Check current size
                buffer = io.BytesIO()
                doc.save(buffer)
                current_size = len(buffer.getvalue())
                attempts += 1
        
        # Final save
        buffer = io.BytesIO()
        doc.save(buffer)
        content_bytes = buffer.getvalue()
        
        return content_bytes


class XLSXTemplate(BaseTemplate):
    """Generate realistic Excel spreadsheets"""
    
    def generate(self, target_size: int, filename: str) -> bytes:
        wb = Workbook()
        ws = wb.active
        
        # Choose spreadsheet type
        sheet_types = ['financial', 'inventory', 'employee', 'sales']
        sheet_type = random.choice(sheet_types)
        
        ws.title = sheet_type.title() + " Data"
        
        # Generate headers and data based on type
        if sheet_type == 'financial':
            headers = ['Date', 'Description', 'Category', 'Amount', 'Balance']
            self._populate_financial_data(ws, headers)
        elif sheet_type == 'inventory':
            headers = ['SKU', 'Product Name', 'Category', 'Quantity', 'Unit Price', 'Total Value']
            self._populate_inventory_data(ws, headers)
        elif sheet_type == 'employee':
            headers = ['Employee ID', 'Name', 'Department', 'Position', 'Salary', 'Start Date']
            self._populate_employee_data(ws, headers)
        else:  # sales
            headers = ['Date', 'Customer', 'Product', 'Quantity', 'Unit Price', 'Total']
            self._populate_sales_data(ws, headers)
        
        # Style the headers
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")
        
        # Add more data if needed to reach approximate target size
        if target_size > 50000:  # Only for larger files
            current_size = 0
            attempts = 0
            while current_size < target_size * 0.8 and attempts < 100:
                self._add_data_row(ws, sheet_type, len(headers))
                
                # Check size every 10 rows to avoid excessive computation
                if attempts % 10 == 0:
                    buffer = io.BytesIO()
                    wb.save(buffer)
                    current_size = len(buffer.getvalue())
                
                attempts += 1
        
        # Final save - don't truncate Excel files as it corrupts them
        buffer = io.BytesIO()
        wb.save(buffer)
        content_bytes = buffer.getvalue()
        
        return content_bytes
    
    def _populate_financial_data(self, ws, headers):
        balance = self.fake.random_int(min=1000, max=50000)
        for row in range(2, random.randint(20, 100)):
            amount = self.fake.random_int(min=-5000, max=5000)
            balance += amount
            ws.cell(row=row, column=1, value=self.fake.date_between(start_date='-1y'))
            ws.cell(row=row, column=2, value=self.fake.company())
            ws.cell(row=row, column=3, value=random.choice(['Income', 'Expense', 'Transfer']))
            ws.cell(row=row, column=4, value=amount)
            ws.cell(row=row, column=5, value=balance)
    
    def _populate_inventory_data(self, ws, headers):
        for row in range(2, random.randint(20, 100)):
            quantity = self.fake.random_int(min=0, max=1000)
            unit_price = self.fake.random_int(min=1, max=500)
            ws.cell(row=row, column=1, value=f"SKU{self.fake.random_int(min=1000, max=9999)}")
            ws.cell(row=row, column=2, value=self.fake.word().title())
            ws.cell(row=row, column=3, value=random.choice(['Electronics', 'Clothing', 'Books', 'Home']))
            ws.cell(row=row, column=4, value=quantity)
            ws.cell(row=row, column=5, value=unit_price)
            ws.cell(row=row, column=6, value=quantity * unit_price)
    
    def _populate_employee_data(self, ws, headers):
        for row in range(2, random.randint(10, 50)):
            ws.cell(row=row, column=1, value=f"EMP{self.fake.random_int(min=1000, max=9999)}")
            ws.cell(row=row, column=2, value=self.fake.name())
            ws.cell(row=row, column=3, value=random.choice(['IT', 'Sales', 'HR', 'Finance', 'Operations']))
            ws.cell(row=row, column=4, value=self.fake.job())
            ws.cell(row=row, column=5, value=self.fake.random_int(min=30000, max=150000))
            ws.cell(row=row, column=6, value=self.fake.date_between(start_date='-5y'))
    
    def _populate_sales_data(self, ws, headers):
        for row in range(2, random.randint(50, 200)):
            quantity = self.fake.random_int(min=1, max=10)
            unit_price = self.fake.random_int(min=10, max=1000)
            ws.cell(row=row, column=1, value=self.fake.date_between(start_date='-6m'))
            ws.cell(row=row, column=2, value=self.fake.company())
            ws.cell(row=row, column=3, value=self.fake.word().title())
            ws.cell(row=row, column=4, value=quantity)
            ws.cell(row=row, column=5, value=unit_price)
            ws.cell(row=row, column=6, value=quantity * unit_price)
    
    def _add_data_row(self, ws, sheet_type, num_cols):
        row = ws.max_row + 1
        if sheet_type == 'financial':
            amount = self.fake.random_int(min=-5000, max=5000)
            ws.cell(row=row, column=1, value=self.fake.date_between(start_date='-1y'))
            ws.cell(row=row, column=2, value=self.fake.company())
            ws.cell(row=row, column=3, value=random.choice(['Income', 'Expense']))
            ws.cell(row=row, column=4, value=amount)
            ws.cell(row=row, column=5, value=self.fake.random_int(min=1000, max=50000))


class EMLTemplate(BaseTemplate):
    """Generate realistic email files"""
    
    def __init__(self, language: str = 'en'):
        super().__init__(language)
        self.image_refs = []  # Store image references for embedding
        self.attachment_refs = []  # Store attachment references
        self.conversation_thread = []  # Store conversation history
    
    def set_image_references(self, image_refs: List[str]):
        """Set list of image filenames to embed in this email"""
        self.image_refs = image_refs
    
    def set_attachment_references(self, attachment_refs: List[str]):
        """Set list of attachment filenames to reference in this email"""
        self.attachment_refs = attachment_refs
    
    def generate(self, target_size: int, filename: str) -> bytes:
        # Generate a single email with proper RFC-822 EML structure
        email_content = self._generate_single_email(target_size)
        
        content_bytes = email_content.encode('utf-8')
        
        # Don't truncate or pad EML files as it corrupts the format
        # The _generate_single_email method handles target size internally
        return content_bytes
    
    def _generate_single_email(self, target_size: int) -> str:
        """Generate a single email with proper RFC-822 EML structure"""
        sender_name = self.fake.name()
        sender_email = self.fake.email()
        recipient_name = self.fake.name()
        recipient_email = self.fake.email()
        subject = self._generate_subject()
        email_date = self.get_random_document_date()
        
        # Generate message ID
        message_id = f"{self.fake.uuid4()}@{sender_email.split('@')[1]}"
        
        # Generate plain text body content
        body_content = self._generate_plain_text_email_body()
        
        # Add extensive additional content to reach target size
        if len(body_content.encode('utf-8')) < target_size * 0.7:
            additional_content = self._generate_extensive_email_content(target_size)
            body_content += f"\n\n{additional_content}"
        
        # Create proper RFC-822 EML structure with headers and body separation
        email_content = f"""From: {sender_name} <{sender_email}>
To: {recipient_name} <{recipient_email}>
Subject: {subject}
Date: {email_date.strftime('%a, %d %b %Y %H:%M:%S +0000')}
Message-ID: <{message_id}>
MIME-Version: 1.0
Content-Type: text/plain; charset=utf-8
Content-Transfer-Encoding: 8bit
X-Priority: Normal
X-Mailer: {random.choice(['Outlook', 'Thunderbird', 'Apple Mail', 'Gmail'])}

{body_content}

--
{sender_name}
{self.fake.job()}
{self.fake.company()}
{self.fake.phone_number()}
"""
        return email_content
    
    def _generate_plain_text_email_body(self) -> str:
        """Generate plain text email body with references"""
        # Choose email type
        email_types = ['business_update', 'project_discussion', 'meeting_request', 'document_review', 'status_report']
        email_type = random.choice(email_types)
        
        # Generate greeting
        greetings = [
            "Dear Team,",
            "Hi everyone,",
            "Hello colleagues,",
            "Good morning,",
            "Dear all,"
        ]
        greeting = random.choice(greetings)
        
        # Generate main content based on type
        if email_type == 'business_update':
            content = self._generate_business_update_text()
        elif email_type == 'project_discussion':
            content = self._generate_project_discussion_text()
        elif email_type == 'meeting_request':
            content = self._generate_meeting_request_text()
        elif email_type == 'document_review':
            content = self._generate_document_review_text()
        else:  # status_report
            content = self._generate_status_report_text()
        
        # Add image references if available
        image_section = ""
        if self.image_refs:
            image_section = "\n\nREFERENCED IMAGES:\n"
            for image_ref in self.image_refs:
                descriptions = [
                    f"- {image_ref}: Chart showing progress metrics",
                    f"- {image_ref}: Visual diagram and flowchart",
                    f"- {image_ref}: Screenshot of current status",
                    f"- {image_ref}: Graphical representation of data"
                ]
                image_section += random.choice(descriptions) + "\n"
        
        # Add attachment references if available
        attachment_section = ""
        if self.attachment_refs:
            attachment_section = "\n\nATTACHED DOCUMENTS:\n"
            for attachment_ref in self.attachment_refs:
                descriptions = [
                    f"- {attachment_ref}: Detailed analysis report",
                    f"- {attachment_ref}: Supporting documentation",
                    f"- {attachment_ref}: Data spreadsheet with calculations",
                    f"- {attachment_ref}: Project specifications and requirements"
                ]
                attachment_section += random.choice(descriptions) + "\n"
        
        # Generate closing
        closings = [
            "Please let me know if you have any questions.",
            "Looking forward to your feedback.",
            "Thank you for your attention to this matter.",
            "Please review and get back to me at your earliest convenience.",
            "I appreciate your time and input on this."
        ]
        closing = random.choice(closings)
        
        # Combine all parts
        full_body = f"""{greeting}

{content}{image_section}{attachment_section}

{closing}

Best regards,"""
        
        return full_body
    
    def _generate_business_update_text(self) -> str:
        """Generate business update email content in plain text"""
        return f"""I hope this email finds you well. I wanted to provide you with an update on our recent business developments and key metrics.

Key Highlights:
- Revenue Growth: We've achieved a {random.randint(5, 25)}% increase in quarterly revenue
- Customer Acquisition: {random.randint(50, 500)} new customers onboarded this month
- Product Development: {random.randint(2, 8)} new features released
- Team Expansion: {random.randint(3, 15)} new team members joined

{self.fake.text(max_nb_chars=200)}

Please review the attached documents for detailed analysis."""
    
    def _generate_project_discussion_text(self) -> str:
        """Generate project discussion email content in plain text"""
        project_name = f"Project {random.choice(['Alpha', 'Beta', 'Gamma', 'Phoenix', 'Titan', 'Nova'])}"
        
        return f"""I wanted to follow up on our discussion regarding {project_name} and share some important updates.

Current Status:
{self.fake.text(max_nb_chars=150)}

Next Steps:
1. {self.fake.sentence()}
2. {self.fake.sentence()}
3. {self.fake.sentence()}

Timeline: We're targeting completion by {self.get_random_date_in_range(days_before=0, days_after=60).strftime('%B %d, %Y')}.

The attached images show our current progress and the documents contain detailed specifications."""
    
    def _generate_meeting_request_text(self) -> str:
        """Generate meeting request email content in plain text"""
        meeting_date = self.get_random_date_in_range(days_before=0, days_after=14)
        
        return f"""I would like to schedule a meeting to discuss our upcoming initiatives and review the current project status.

Meeting Details:
- Date: {meeting_date.strftime('%A, %B %d, %Y')}
- Time: {random.randint(9, 16)}:00 - {random.randint(10, 17)}:00
- Location: {random.choice(['Conference Room A', 'Main Boardroom', 'Virtual (Zoom)', 'Office 205'])}

Agenda:
1. Review of current metrics and KPIs
2. Discussion of upcoming deadlines
3. Resource allocation and budget review
4. Q&A session

Please review the attached materials before the meeting. The images provide visual context for our discussion points.

Please confirm your attendance by replying to this email."""
    
    def _generate_document_review_text(self) -> str:
        """Generate document review email content in plain text"""
        return f"""I've completed the initial review of the documents we discussed and wanted to share my feedback with you.

Review Summary:
{self.fake.text(max_nb_chars=200)}

Key Observations:
- Strengths: {self.fake.sentence()}
- Areas for improvement: {self.fake.sentence()}
- Recommendations: {self.fake.sentence()}

I've attached the reviewed documents with my comments and suggestions. The images referenced in the documents provide additional context for the proposed changes.

Please take a look and let me know your thoughts. We should schedule a follow-up meeting to discuss the next steps."""
    
    def _generate_status_report_text(self) -> str:
        """Generate status report email content in plain text"""
        return f"""Please find below the weekly status report for our ongoing projects and initiatives.

Project Status Overview:
- Initiative A: On Track ({random.randint(60, 95)}% complete)
- Initiative B: At Risk ({random.randint(30, 70)}% complete)
- Initiative C: Completed (100% complete)

Key Metrics:
{self.fake.text(max_nb_chars=150)}

Detailed charts and analysis are available in the attached documents. Visual progress indicators are shown in the referenced images.

Please don't hesitate to reach out if you need any clarification or additional information."""
    
    def _generate_extensive_email_content(self, target_size: int) -> str:
        """Generate extensive additional email content to reach target size"""
        content_parts = []
        
        # Add detailed business content
        content_parts.append("\nAdditional Information:")
        
        # Generate multiple sections of realistic business content
        sections = [
            "Style report age. Stand fact arm shake happy itself step over. Laugh life check generation large commercial movement serious.",
            "Reduce present mind book. Pull board improve raise feeling blue mouth.",
            "Professor individual movie Congress safe. A agree memory score stuff site. Wish industry court position price owner though.",
            "Join establish break third particular hope sort. Voice must impact feel assume.",
            "Security rate investment turn design. Specific art see. Professional allow upon maybe teacher shoulder question. Region usually sure opportunity material west."
        ]
        
        # Add extensive realistic content using phrase generator or faker
        for i in range(min(50, target_size // 1000)):  # Scale content based on target size
            if self.phrase_gen:
                content_parts.append(self.phrase_gen.generate_realistic_document_content(200))
            else:
                content_parts.append(self.fake.text(max_nb_chars=200))
            
            # Add some business-like structure occasionally
            if i % 5 == 0:
                content_parts.append(f"\n{random.choice(['Analysis:', 'Summary:', 'Details:', 'Notes:', 'Review:'])}")
                content_parts.append(random.choice(sections))
        
        # Join all content
        full_content = '\n'.join(content_parts)
        
        # Ensure we have enough content
        while len(full_content.encode('utf-8')) < target_size * 0.6:
            if self.phrase_gen:
                additional = self.phrase_gen.generate_realistic_document_content(500)
            else:
                additional = self.fake.text(max_nb_chars=500)
            full_content += f"\n{additional}"
            
            # Prevent infinite loop
            if len(full_content) > target_size * 2:
                break
        
        return full_content
    
    def _generate_conversation_thread(self, thread_length: int, target_size: int) -> str:
        """Generate a conversation thread with multiple emails"""
        
        # Create participants
        participants = []
        import random as rand
        for _ in range(rand.randint(2, 4)):  # 2-4 people in conversation
            participants.append({
                'name': self.fake.name(),
                'email': self.fake.email(),
                'title': self.fake.job(),
                'company': self.fake.company()
            })
        
        # Generate base subject
        base_subject = self._generate_subject()
        
        # Generate conversation
        conversation = []
        current_date = self.get_random_document_date()
        
        for i in range(thread_length):
            # Choose sender and recipients
            sender = rand.choice(participants)
            recipients = [p for p in participants if p != sender]
            
            # Determine subject (Re: for replies)
            if i == 0:
                subject = base_subject
            else:
                subject = f"Re: {base_subject}"
                if i > 2:  # Add more Re: for longer threads
                    subject = f"Re: {subject}"
            
            # Generate email date (chronological order)
            if i > 0:
                # Add random time between emails (minutes to hours)
                from datetime import timedelta
                time_delta = timedelta(
                    minutes=rand.randint(5, 180),  # 5 minutes to 3 hours
                    hours=rand.randint(0, 24) if rand.random() < 0.3 else 0  # Sometimes days apart
                )
                current_date = current_date + time_delta
            
            # Generate email body based on position in thread
            if i == 0:
                body = self._generate_thread_starter()
            elif i == thread_length - 1:
                body = self._generate_thread_closer()
            else:
                body = self._generate_thread_reply(i)
            
            # Add image/attachment references to some emails
            if self.image_refs and rand.random() < 0.4:  # 40% chance
                body += self._add_image_references_to_body()
            
            if self.attachment_refs and rand.random() < 0.5:  # 50% chance
                body += self._add_attachment_references_to_body()
            
            # Create email
            email = f"""From: {sender['name']} <{sender['email']}>
To: {', '.join([f"{r['name']} <{r['email']}>" for r in recipients[:2]])}
Subject: {subject}
Date: {current_date.strftime('%a, %d %b %Y %H:%M:%S +0000')}
Message-ID: <{self.fake.uuid4()}@{sender['email'].split('@')[1]}>
MIME-Version: 1.0
Content-Type: text/html; charset=utf-8
Content-Transfer-Encoding: 8bit
X-Priority: Normal
In-Reply-To: <{self.fake.uuid4()}@{participants[0]['email'].split('@')[1]}>
References: <{self.fake.uuid4()}@{participants[0]['email'].split('@')[1]}>

{body}

--
{sender['name']}
{sender['title']}
{sender['company']}
{self.fake.phone_number()}

"""
            conversation.append(email)
        
        # Combine all emails in the thread (newest first, like most email clients)
        return '\n'.join(reversed(conversation))
    
    def _generate_additional_thread_emails(self, padding_needed: int) -> str:
        """Generate additional emails for padding"""
        additional_emails = []
        emails_to_add = min(3, padding_needed // 500)  # Rough estimate
        
        for i in range(emails_to_add):
            sender_name = self.fake.name()
            sender_email = self.fake.email()
            subject = f"Re: {self._generate_subject()}"
            email_date = self.get_random_document_date()
            body = self._generate_thread_reply(i + 10)  # Different reply type
            
            email = f"""
From: {sender_name} <{sender_email}>
To: {self.fake.name()} <{self.fake.email()}>
Subject: {subject}
Date: {email_date.strftime('%a, %d %b %Y %H:%M:%S +0000')}
Message-ID: <{self.fake.uuid4()}@{sender_email.split('@')[1]}>

{body}

--
{sender_name}
{self.fake.job()}
"""
            additional_emails.append(email)
        
        return '\n'.join(additional_emails)
    
    def _generate_subject(self) -> str:
        subjects = [
            f"RE: {self.fake.catch_phrase()}",
            f"Meeting Request: {self.fake.date()}",
            f"Project Update - {self.fake.word().title()}",
            f"Invoice #{self.fake.random_int(min=1000, max=9999)}",
            f"Important: {self.fake.bs()}",
            f"FW: {self.fake.company()} Proposal",
            f"Quarterly Report - {self.fake.date()}",
            f"Action Required: {self.fake.word().title()}"
        ]
        return random.choice(subjects)
    
    def _generate_rich_email_body(self) -> str:
        """Generate rich HTML email body with images and attachments"""
        
        # Choose email type
        email_types = ['business_update', 'project_discussion', 'meeting_request', 'document_review', 'status_report']
        email_type = random.choice(email_types)
        
        # Generate content based on type
        if email_type == 'business_update':
            content = self._generate_business_update_content()
        elif email_type == 'project_discussion':
            content = self._generate_project_discussion_content()
        elif email_type == 'meeting_request':
            content = self._generate_meeting_request_content()
        elif email_type == 'document_review':
            content = self._generate_document_review_content()
        else:  # status_report
            content = self._generate_status_report_content()
        
        # Wrap in HTML
        html_body = f"""<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
        .header {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        .content {{ margin: 20px 0; }}
        .image-ref {{ background-color: #f8f9fa; padding: 10px; margin: 10px 0; border-left: 4px solid #007bff; }}
        .attachment-ref {{ background-color: #fff3cd; padding: 10px; margin: 10px 0; border-left: 4px solid #ffc107; }}
        .footer {{ color: #6c757d; font-size: 0.9em; margin-top: 30px; }}
    </style>
</head>
<body>
    <div class="header">
        <h2>{self._get_email_title(email_type)}</h2>
    </div>
    
    <div class="content">
        {content}
    </div>
    
    {self._add_image_references_html() if self.image_refs else ''}
    {self._add_attachment_references_html() if self.attachment_refs else ''}
    
    <div class="footer">
        <p>This email was sent on {self.get_random_document_date().strftime('%B %d, %Y')}.</p>
        <p>Please consider the environment before printing this email.</p>
    </div>
</body>
</html>"""
        
        return html_body
    
    def _generate_business_update_content(self) -> str:
        """Generate business update email content"""
        return f"""
        <p>Dear Team,</p>
        
        <p>I hope this email finds you well. I wanted to provide you with an update on our recent business developments and key metrics.</p>
        
        <h3>Key Highlights:</h3>
        <ul>
            <li><strong>Revenue Growth:</strong> We've achieved a {random.randint(5, 25)}% increase in quarterly revenue</li>
            <li><strong>Customer Acquisition:</strong> {random.randint(50, 500)} new customers onboarded this month</li>
            <li><strong>Product Development:</strong> {random.randint(2, 8)} new features released</li>
            <li><strong>Team Expansion:</strong> {random.randint(3, 15)} new team members joined</li>
        </ul>
        
        <p>{self.fake.text(max_nb_chars=300)}</p>
        
        <p>Please review the attached documents for detailed analysis and let me know if you have any questions.</p>
        
        <p>Best regards,</p>
        """
    
    def _generate_project_discussion_content(self) -> str:
        """Generate project discussion email content"""
        project_name = f"Project {random.choice(['Alpha', 'Beta', 'Gamma', 'Phoenix', 'Titan', 'Nova'])}"
        
        return f"""
        <p>Hi everyone,</p>
        
        <p>I wanted to follow up on our discussion regarding <strong>{project_name}</strong> and share some important updates.</p>
        
        <h3>Current Status:</h3>
        <p>{self.fake.text(max_nb_chars=200)}</p>
        
        <h3>Next Steps:</h3>
        <ol>
            <li>{self.fake.sentence()}</li>
            <li>{self.fake.sentence()}</li>
            <li>{self.fake.sentence()}</li>
        </ol>
        
        <p><strong>Timeline:</strong> We're targeting completion by {self.get_random_date_in_range(days_before=0, days_after=60).strftime('%B %d, %Y')}.</p>
        
        <p>The attached images show our current progress and the documents contain detailed specifications.</p>
        
        <p>Looking forward to your feedback.</p>
        """
    
    def _generate_meeting_request_content(self) -> str:
        """Generate meeting request email content"""
        meeting_date = self.get_random_date_in_range(days_before=0, days_after=14)
        
        return f"""
        <p>Dear colleagues,</p>
        
        <p>I would like to schedule a meeting to discuss our upcoming initiatives and review the current project status.</p>
        
        <h3>Meeting Details:</h3>
        <ul>
            <li><strong>Date:</strong> {meeting_date.strftime('%A, %B %d, %Y')}</li>
            <li><strong>Time:</strong> {random.randint(9, 16)}:00 - {random.randint(10, 17)}:00</li>
            <li><strong>Location:</strong> {random.choice(['Conference Room A', 'Main Boardroom', 'Virtual (Zoom)', 'Office 205'])}</li>
        </ul>
        
        <h3>Agenda:</h3>
        <ol>
            <li>Review of current metrics and KPIs</li>
            <li>Discussion of upcoming deadlines</li>
            <li>Resource allocation and budget review</li>
            <li>Q&A session</li>
        </ol>
        
        <p>Please review the attached materials before the meeting. The images provide visual context for our discussion points.</p>
        
        <p>Please confirm your attendance by replying to this email.</p>
        """
    
    def _generate_document_review_content(self) -> str:
        """Generate document review email content"""
        return f"""
        <p>Hello team,</p>
        
        <p>I've completed the initial review of the documents we discussed and wanted to share my feedback with you.</p>
        
        <h3>Review Summary:</h3>
        <p>{self.fake.text(max_nb_chars=250)}</p>
        
        <h3>Key Observations:</h3>
        <ul>
            <li><strong>Strengths:</strong> {self.fake.sentence()}</li>
            <li><strong>Areas for improvement:</strong> {self.fake.sentence()}</li>
            <li><strong>Recommendations:</strong> {self.fake.sentence()}</li>
        </ul>
        
        <p>I've attached the reviewed documents with my comments and suggestions. The images referenced in the documents provide additional context for the proposed changes.</p>
        
        <p>Please take a look and let me know your thoughts. We should schedule a follow-up meeting to discuss the next steps.</p>
        
        <p>Thank you for your time and attention to this matter.</p>
        """
    
    def _generate_status_report_content(self) -> str:
        """Generate status report email content"""
        return f"""
        <p>Dear stakeholders,</p>
        
        <p>Please find below the weekly status report for our ongoing projects and initiatives.</p>
        
        <h3>Project Status Overview:</h3>
        <table border="1" style="border-collapse: collapse; width: 100%;">
            <tr style="background-color: #f2f2f2;">
                <th style="padding: 8px;">Project</th>
                <th style="padding: 8px;">Status</th>
                <th style="padding: 8px;">Progress</th>
            </tr>
            <tr>
                <td style="padding: 8px;">Initiative A</td>
                <td style="padding: 8px; color: green;">On Track</td>
                <td style="padding: 8px;">{random.randint(60, 95)}%</td>
            </tr>
            <tr>
                <td style="padding: 8px;">Initiative B</td>
                <td style="padding: 8px; color: orange;">At Risk</td>
                <td style="padding: 8px;">{random.randint(30, 70)}%</td>
            </tr>
            <tr>
                <td style="padding: 8px;">Initiative C</td>
                <td style="padding: 8px; color: green;">Completed</td>
                <td style="padding: 8px;">100%</td>
            </tr>
        </table>
        
        <h3>Key Metrics:</h3>
        <p>{self.fake.text(max_nb_chars=200)}</p>
        
        <p>Detailed charts and analysis are available in the attached documents. Visual progress indicators are shown in the referenced images.</p>
        
        <p>Please don't hesitate to reach out if you need any clarification or additional information.</p>
        """
    
    def _get_email_title(self, email_type: str) -> str:
        """Get appropriate title for email type"""
        titles = {
            'business_update': 'Quarterly Business Update',
            'project_discussion': 'Project Status Discussion',
            'meeting_request': 'Meeting Request - Team Sync',
            'document_review': 'Document Review and Feedback',
            'status_report': 'Weekly Status Report'
        }
        return titles.get(email_type, 'Important Update')
    
    def _add_image_references_html(self) -> str:
        """Add HTML section for image references"""
        if not self.image_refs:
            return ""
        
        html = '<div class="image-ref"><h4>📷 Referenced Images:</h4><ul>'
        for image_ref in self.image_refs:
            descriptions = [
                f"Chart showing progress metrics - {image_ref}",
                f"Visual diagram and flowchart - {image_ref}",
                f"Screenshot of current status - {image_ref}",
                f"Graphical representation - {image_ref}",
                f"Process visualization - {image_ref}"
            ]
            html += f'<li>{random.choice(descriptions)}</li>'
        html += '</ul></div>'
        return html
    
    def _add_attachment_references_html(self) -> str:
        """Add HTML section for attachment references"""
        if not self.attachment_refs:
            return ""
        
        html = '<div class="attachment-ref"><h4>📎 Attached Documents:</h4><ul>'
        for attachment_ref in self.attachment_refs:
            descriptions = [
                f"Detailed analysis report - {attachment_ref}",
                f"Supporting documentation - {attachment_ref}",
                f"Data spreadsheet - {attachment_ref}",
                f"Project specifications - {attachment_ref}",
                f"Meeting notes and action items - {attachment_ref}"
            ]
            html += f'<li>{random.choice(descriptions)}</li>'
        html += '</ul></div>'
        return html
    
    def _add_image_references_to_body(self) -> str:
        """Add image references to plain text body"""
        if not self.image_refs:
            return ""
        
        text = "\n\nREFERENCED IMAGES:\n"
        for image_ref in self.image_refs:
            text += f"- {image_ref}: Visual documentation and charts\n"
        return text
    
    def _add_attachment_references_to_body(self) -> str:
        """Add attachment references to plain text body"""
        if not self.attachment_refs:
            return ""
        
        text = "\n\nATTACHED DOCUMENTS:\n"
        for attachment_ref in self.attachment_refs:
            text += f"- {attachment_ref}: Supporting documentation\n"
        return text
    
    def _generate_thread_starter(self) -> str:
        """Generate the first email in a conversation thread"""
        starters = [
            f"""<p>Hi everyone,</p>
            
            <p>I wanted to start a discussion about {self.fake.catch_phrase().lower()} and get everyone's input on the best approach.</p>
            
            <p>{self.fake.text(max_nb_chars=200)}</p>
            
            <p>What are your thoughts on this? I'd love to hear your perspectives.</p>""",
            
            f"""<p>Dear team,</p>
            
            <p>Following up on our conversation about the {self.fake.word()} project. I think we need to address a few key points:</p>
            
            <ul>
                <li>{self.fake.sentence()}</li>
                <li>{self.fake.sentence()}</li>
                <li>{self.fake.sentence()}</li>
            </ul>
            
            <p>Looking forward to your feedback.</p>""",
            
            f"""<p>Hello all,</p>
            
            <p>I hope you're all doing well. I wanted to bring up an important topic that requires our collective attention.</p>
            
            <p>{self.fake.text(max_nb_chars=300)}</p>
            
            <p>Please share your thoughts when you have a chance.</p>"""
        ]
        return random.choice(starters)
    
    def _generate_thread_reply(self, position: int) -> str:
        """Generate a reply email in a conversation thread"""
        replies = [
            f"""<p>Thanks for bringing this up!</p>
            
            <p>I agree with your points, especially regarding {self.fake.word()}. Here are my additional thoughts:</p>
            
            <p>{self.fake.text(max_nb_chars=150)}</p>
            
            <p>What does everyone else think?</p>""",
            
            f"""<p>Great discussion so far.</p>
            
            <p>I'd like to add that we should also consider {self.fake.sentence().lower()}</p>
            
            <p>This could help us {self.fake.sentence().lower()}</p>
            
            <p>Thoughts?</p>""",
            
            f"""<p>I have some concerns about this approach.</p>
            
            <p>{self.fake.text(max_nb_chars=200)}</p>
            
            <p>Perhaps we should explore alternative solutions?</p>""",
            
            f"""<p>Excellent points everyone!</p>
            
            <p>Based on the discussion so far, I think we're moving in the right direction. Let me summarize what I'm hearing:</p>
            
            <p>{self.fake.text(max_nb_chars=180)}</p>
            
            <p>Does this capture everyone's input accurately?</p>"""
        ]
        return random.choice(replies)
    
    def _generate_thread_closer(self) -> str:
        """Generate the final email in a conversation thread"""
        closers = [
            f"""<p>Thank you all for the valuable input!</p>
            
            <p>Based on our discussion, I think we have a clear path forward. I'll summarize the key decisions and next steps:</p>
            
            <p>{self.fake.text(max_nb_chars=200)}</p>
            
            <p>I'll send out a follow-up with action items by end of week.</p>
            
            <p>Thanks again for your time and insights!</p>""",
            
            f"""<p>This has been a very productive discussion.</p>
            
            <p>I appreciate everyone's contributions and perspectives. Let's move forward with the plan we've outlined.</p>
            
            <p>I'll coordinate the next steps and keep everyone updated on progress.</p>
            
            <p>Have a great rest of your week!</p>""",
            
            f"""<p>Perfect! I think we're all aligned now.</p>
            
            <p>{self.fake.text(max_nb_chars=150)}</p>
            
            <p>Let's schedule a follow-up meeting next week to review progress.</p>
            
            <p>Thanks everyone!</p>"""
        ]
        return random.choice(closers)
    
    def _generate_email_body(self) -> str:
        """Legacy method for backward compatibility"""
        return self._generate_rich_email_body()


class CSVTemplate(BaseTemplate):
    """Generate realistic CSV files"""
    
    def generate(self, target_size: int, filename: str) -> bytes:
        buffer = io.StringIO()
        
        # Enhanced CSV types with more realistic business data
        csv_types = ['customers', 'products', 'transactions', 'employees', 'sales_leads', 'inventory', 'orders', 'vendors']
        csv_type = random.choice(csv_types)
        
        if csv_type == 'customers':
            headers = ['Customer_ID', 'First_Name', 'Last_Name', 'Email', 'Phone', 'Address', 'City', 'State', 'Zip', 'Country', 'Registration_Date', 'Last_Purchase', 'Total_Spent', 'Customer_Type']
            self._write_customer_data(buffer, headers)
        elif csv_type == 'products':
            headers = ['Product_ID', 'SKU', 'Name', 'Category', 'Subcategory', 'Price', 'Cost', 'Stock', 'Supplier', 'Description', 'Status', 'Date_Added']
            self._write_product_data(buffer, headers)
        elif csv_type == 'transactions':
            headers = ['Transaction_ID', 'Date', 'Customer_ID', 'Product_ID', 'Quantity', 'Unit_Price', 'Total_Amount', 'Tax', 'Payment_Method', 'Status', 'Sales_Rep']
            self._write_transaction_data(buffer, headers)
        elif csv_type == 'employees':
            headers = ['Employee_ID', 'First_Name', 'Last_Name', 'Email', 'Department', 'Position', 'Salary', 'Hire_Date', 'Manager', 'Location', 'Status']
            self._write_employee_data(buffer, headers)
        elif csv_type == 'sales_leads':
            headers = ['Lead_ID', 'Company', 'Contact_Name', 'Email', 'Phone', 'Industry', 'Lead_Source', 'Status', 'Score', 'Assigned_To', 'Date_Created']
            self._write_sales_leads_data(buffer, headers)
        elif csv_type == 'inventory':
            headers = ['Item_ID', 'Name', 'Location', 'Quantity', 'Min_Stock', 'Max_Stock', 'Reorder_Point', 'Last_Restock', 'Supplier', 'Unit_Cost']
            self._write_inventory_data(buffer, headers)
        elif csv_type == 'orders':
            headers = ['Order_ID', 'Customer_ID', 'Order_Date', 'Ship_Date', 'Total_Amount', 'Shipping_Cost', 'Status', 'Payment_Status', 'Shipping_Address']
            self._write_orders_data(buffer, headers)
        else:  # vendors
            headers = ['Vendor_ID', 'Company_Name', 'Contact_Name', 'Email', 'Phone', 'Address', 'City', 'State', 'Payment_Terms', 'Status']
            self._write_vendors_data(buffer, headers)
        
        content = buffer.getvalue()
        content_bytes = content.encode('utf-8')
        
        # Add more rows if needed to reach approximate target size
        if len(content_bytes) < target_size and target_size > 10000:
            target_rows = min((target_size - len(content_bytes)) // 200, 5000)  # More realistic estimate
            for _ in range(target_rows):
                buffer.write('\n')
                self._write_data_row(buffer, csv_type)
                
                # Check size periodically
                if _ % 100 == 0:
                    current_size = len(buffer.getvalue().encode('utf-8'))
                    if current_size >= target_size * 0.9:
                        break
            
            content_bytes = buffer.getvalue().encode('utf-8')
        
        return content_bytes
    
    def _write_customer_data(self, buffer, headers):
        writer = csv.writer(buffer)
        writer.writerow(headers)
        
        customer_types = ['Premium', 'Standard', 'Basic', 'VIP', 'Corporate']
        
        for i in range(random.randint(50, 200)):
            if self.phrase_gen:
                name_data = self.phrase_gen.generate_realistic_name_data()
                first_name = name_data['name'].split()[0] if ' ' in name_data['name'] else name_data['name']
                last_name = name_data['name'].split()[-1] if ' ' in name_data['name'] else self.fake.last_name()
                # Use randomized dates for customer registration and last purchase
                reg_date = self.get_random_date_in_range(days_before=730, days_after=0)  # Up to 2 years ago
                last_purchase = self.get_random_date_in_range(days_before=180, days_after=0)  # Up to 6 months ago
                
                writer.writerow([
                    f"CUST{i+1000}",
                    first_name,
                    last_name,
                    self.fake.email(),
                    self.fake.phone_number(),
                    name_data['address'],
                    name_data['city'],
                    name_data['state'],
                    self.fake.zipcode(),
                    'USA',
                    reg_date.strftime('%Y-%m-%d'),
                    last_purchase.strftime('%Y-%m-%d'),
                    f"${self.fake.random_int(min=100, max=50000):.2f}",
                    random.choice(customer_types)
                ])
            else:
                # Use randomized dates for customer registration and last purchase
                reg_date = self.get_random_date_in_range(days_before=730, days_after=0)  # Up to 2 years ago
                last_purchase = self.get_random_date_in_range(days_before=180, days_after=0)  # Up to 6 months ago
                
                writer.writerow([
                    f"CUST{i+1000}",
                    self.fake.first_name(),
                    self.fake.last_name(),
                    self.fake.email(),
                    self.fake.phone_number(),
                    self.fake.street_address(),
                    self.fake.city(),
                    self.fake.state(),
                    self.fake.zipcode(),
                    'USA',
                    reg_date.strftime('%Y-%m-%d'),
                    last_purchase.strftime('%Y-%m-%d'),
                    f"${self.fake.random_int(min=100, max=50000):.2f}",
                    random.choice(customer_types)
                ])
    
    def _write_product_data(self, buffer, headers):
        writer = csv.writer(buffer)
        writer.writerow(headers)
        
        categories = ['Electronics', 'Clothing', 'Books', 'Home & Garden', 'Sports', 'Toys']
        
        for i in range(random.randint(100, 1000)):
            writer.writerow([
                f"PROD{i+1000}",
                self.fake.word().title(),
                random.choice(categories),
                f"{self.fake.random_int(min=5, max=1000):.2f}",
                self.fake.random_int(min=0, max=500),
                self.fake.company(),
                self.fake.text(max_nb_chars=100)
            ])
    
    def _write_transaction_data(self, buffer, headers):
        writer = csv.writer(buffer)
        writer.writerow(headers)
        
        payment_methods = ['Credit Card', 'Debit Card', 'PayPal', 'Cash', 'Check']
        
        for i in range(random.randint(200, 2000)):
            quantity = self.fake.random_int(min=1, max=10)
            price = self.fake.random_int(min=10, max=500)
            writer.writerow([
                f"TXN{i+10000}",
                self.fake.date_between(start_date='-1y').strftime('%Y-%m-%d'),
                f"CUST{self.fake.random_int(min=1000, max=1500)}",
                f"PROD{self.fake.random_int(min=1000, max=1100)}",
                quantity,
                f"{quantity * price:.2f}",
                random.choice(payment_methods)
            ])
    
    def _write_employee_data(self, buffer, headers):
        writer = csv.writer(buffer)
        writer.writerow(headers)
        
        departments = ['IT', 'Sales', 'HR', 'Finance', 'Operations', 'Marketing']
        
        for i in range(random.randint(20, 200)):
            writer.writerow([
                f"EMP{i+1000}",
                self.fake.name(),
                random.choice(departments),
                self.fake.job(),
                self.fake.random_int(min=30000, max=150000),
                self.fake.date_between(start_date='-5y').strftime('%Y-%m-%d'),
                self.fake.name()
            ])
    
    def _write_sales_leads_data(self, buffer, headers):
        writer = csv.writer(buffer)
        writer.writerow(headers)
        
        industries = ['Technology', 'Healthcare', 'Finance', 'Manufacturing', 'Retail', 'Real Estate', 'Education']
        lead_sources = ['Website', 'Cold Call', 'Email Campaign', 'Referral', 'Trade Show', 'Social Media']
        statuses = ['New', 'Qualified', 'Proposal', 'Negotiation', 'Closed-Won', 'Closed-Lost']
        
        for i in range(random.randint(20, 100)):
            writer.writerow([
                f"LEAD{i+2000}",
                self.fake.company(),
                self.fake.name(),
                self.fake.email(),
                self.fake.phone_number(),
                random.choice(industries),
                random.choice(lead_sources),
                random.choice(statuses),
                random.randint(1, 100),
                self.fake.name(),
                self.fake.date_between(start_date='-6m').strftime('%Y-%m-%d')
            ])
    
    def _write_inventory_data(self, buffer, headers):
        writer = csv.writer(buffer)
        writer.writerow(headers)
        
        locations = ['Warehouse A', 'Warehouse B', 'Store Front', 'Distribution Center', 'Back Room']
        
        for i in range(random.randint(50, 300)):
            quantity = random.randint(0, 1000)
            min_stock = random.randint(10, 50)
            max_stock = min_stock + random.randint(100, 500)
            reorder_point = random.randint(min_stock, min_stock + 20)
            
            writer.writerow([
                f"ITEM{i+3000}",
                self.fake.word().title() + " " + self.fake.word().title(),
                random.choice(locations),
                quantity,
                min_stock,
                max_stock,
                reorder_point,
                self.fake.date_between(start_date='-3m').strftime('%Y-%m-%d'),
                self.fake.company(),
                f"${self.fake.random_int(min=5, max=500):.2f}"
            ])
    
    def _write_orders_data(self, buffer, headers):
        writer = csv.writer(buffer)
        writer.writerow(headers)
        
        statuses = ['Pending', 'Processing', 'Shipped', 'Delivered', 'Cancelled']
        payment_statuses = ['Paid', 'Pending', 'Failed', 'Refunded']
        
        for i in range(random.randint(30, 200)):
            order_date = self.fake.date_between(start_date='-6m')
            ship_date = self.fake.date_between(start_date=order_date, end_date='today')
            total_amount = self.fake.random_int(min=50, max=5000)
            shipping_cost = random.uniform(5.99, 29.99)
            
            writer.writerow([
                f"ORD{i+4000}",
                f"CUST{random.randint(1000, 1500)}",
                order_date.strftime('%Y-%m-%d'),
                ship_date.strftime('%Y-%m-%d'),
                f"${total_amount:.2f}",
                f"${shipping_cost:.2f}",
                random.choice(statuses),
                random.choice(payment_statuses),
                self.fake.address().replace('\n', ', ')
            ])
    
    def _write_vendors_data(self, buffer, headers):
        writer = csv.writer(buffer)
        writer.writerow(headers)
        
        payment_terms = ['Net 30', 'Net 60', '2/10 Net 30', 'Cash on Delivery', 'Prepaid']
        statuses = ['Active', 'Inactive', 'Pending', 'Suspended']
        
        for i in range(random.randint(10, 50)):
            writer.writerow([
                f"VEND{i+5000}",
                self.fake.company(),
                self.fake.name(),
                self.fake.email(),
                self.fake.phone_number(),
                self.fake.street_address(),
                self.fake.city(),
                self.fake.state(),
                random.choice(payment_terms),
                random.choice(statuses)
            ])
    
    def _write_data_row(self, buffer, csv_type):
        writer = csv.writer(buffer)
        
        if csv_type == 'customers':
            customer_types = ['Premium', 'Standard', 'Basic', 'VIP', 'Corporate']
            writer.writerow([
                f"CUST{self.fake.random_int(min=1000, max=9999)}",
                self.fake.first_name(),
                self.fake.last_name(),
                self.fake.email(),
                self.fake.phone_number(),
                self.fake.street_address(),
                self.fake.city(),
                self.fake.state(),
                self.fake.zipcode(),
                'USA',
                self.fake.date_between(start_date='-2y').strftime('%Y-%m-%d'),
                self.fake.date_between(start_date='-6m').strftime('%Y-%m-%d'),
                f"${self.fake.random_int(min=100, max=50000):.2f}",
                random.choice(customer_types)
            ])
        elif csv_type == 'sales_leads':
            industries = ['Technology', 'Healthcare', 'Finance', 'Manufacturing', 'Retail']
            lead_sources = ['Website', 'Cold Call', 'Email Campaign', 'Referral']
            statuses = ['New', 'Qualified', 'Proposal', 'Closed-Won', 'Closed-Lost']
            writer.writerow([
                f"LEAD{self.fake.random_int(min=2000, max=2999)}",
                self.fake.company(),
                self.fake.name(),
                self.fake.email(),
                self.fake.phone_number(),
                random.choice(industries),
                random.choice(lead_sources),
                random.choice(statuses),
                random.randint(1, 100),
                self.fake.name(),
                self.fake.date_between(start_date='-6m').strftime('%Y-%m-%d')
            ])
        elif csv_type == 'inventory':
            locations = ['Warehouse A', 'Warehouse B', 'Store Front']
            writer.writerow([
                f"ITEM{self.fake.random_int(min=3000, max=3999)}",
                self.fake.word().title(),
                random.choice(locations),
                random.randint(0, 1000),
                random.randint(10, 50),
                random.randint(100, 500),
                random.randint(20, 80),
                self.fake.date_between(start_date='-3m').strftime('%Y-%m-%d'),
                self.fake.company(),
                f"${self.fake.random_int(min=5, max=500):.2f}"
            ])
        elif csv_type == 'orders':
            statuses = ['Pending', 'Processing', 'Shipped', 'Delivered']
            payment_statuses = ['Paid', 'Pending', 'Failed']
            writer.writerow([
                f"ORD{self.fake.random_int(min=4000, max=4999)}",
                f"CUST{random.randint(1000, 1500)}",
                self.fake.date_between(start_date='-6m').strftime('%Y-%m-%d'),
                self.fake.date_between(start_date='-3m').strftime('%Y-%m-%d'),
                f"${self.fake.random_int(min=50, max=5000):.2f}",
                f"${random.uniform(5.99, 29.99):.2f}",
                random.choice(statuses),
                random.choice(payment_statuses),
                self.fake.address().replace('\n', ', ')
            ])
        elif csv_type == 'vendors':
            payment_terms = ['Net 30', 'Net 60', '2/10 Net 30']
            statuses = ['Active', 'Inactive', 'Pending']
            writer.writerow([
                f"VEND{self.fake.random_int(min=5000, max=5999)}",
                self.fake.company(),
                self.fake.name(),
                self.fake.email(),
                self.fake.phone_number(),
                self.fake.street_address(),
                self.fake.city(),
                self.fake.state(),
                random.choice(payment_terms),
                random.choice(statuses)
            ])


class TXTTemplate(BaseTemplate):
    """Generate realistic text files"""
    
    def generate(self, target_size: int, filename: str) -> bytes:
        # Choose text type
        text_types = ['notes', 'log', 'documentation', 'configuration']
        text_type = random.choice(text_types)
        
        if text_type == 'notes':
            content = self._generate_notes()
        elif text_type == 'log':
            content = self._generate_log()
        elif text_type == 'documentation':
            content = self._generate_documentation()
        else:  # configuration
            content = self._generate_configuration()
        
        content_bytes = content.encode('utf-8')
        
        # Repeat content to reach target size
        if len(content_bytes) < target_size:
            repetitions = (target_size // len(content_bytes)) + 1
            content_bytes = (content_bytes * repetitions)[:target_size]
        
        return content_bytes
    
    def _generate_notes(self) -> str:
        if self.phrase_gen:
            content = self.phrase_gen.generate_realistic_document_content(500)
            name_data1 = self.phrase_gen.generate_realistic_name_data()
            name_data2 = self.phrase_gen.generate_realistic_name_data()
            
            todos = [
                self.phrase_gen.generate_query(),
                self.phrase_gen.generate_sentence(),
                self.phrase_gen.generate_query()
            ]
            
            note_date = self.get_random_document_date()
            return f"""Personal Notes
================

Date: {note_date.strftime('%B %d, %Y')}

{content}

TODO:
- {todos[0]}
- {todos[1]}
- {todos[2]}

Ideas:
{self.phrase_gen.generate_realistic_document_content(300)}

Important Contacts:
- {name_data1['name']}: {self.fake.phone_number()}
- {name_data2['name']}: {self.fake.email()}
"""
        else:
            note_date = self.get_random_document_date()
            return f"""Personal Notes
================

Date: {note_date.strftime('%B %d, %Y')}

{self.fake.text(max_nb_chars=500)}

TODO:
- {self.fake.sentence()}
- {self.fake.sentence()}
- {self.fake.sentence()}

Ideas:
{self.fake.text(max_nb_chars=300)}

Important Contacts:
- {self.fake.name()}: {self.fake.phone_number()}
- {self.fake.name()}: {self.fake.email()}
"""
    
    def _generate_log(self) -> str:
        entries = []
        for _ in range(random.randint(10, 50)):
            timestamp = self.fake.date_time_between(start_date='-30d')
            level = random.choice(['INFO', 'WARNING', 'ERROR', 'DEBUG'])
            message = self.fake.sentence()
            entries.append(f"[{timestamp.strftime('%Y-%m-%d %H:%M:%S')}] {level}: {message}")
        
        return f"""System Log File
===============

{chr(10).join(entries)}
"""
    
    def _generate_documentation(self) -> str:
        return f"""Project Documentation
=====================

Overview:
{self.fake.text(max_nb_chars=400)}

Installation:
1. {self.fake.sentence()}
2. {self.fake.sentence()}
3. {self.fake.sentence()}

Configuration:
{self.fake.text(max_nb_chars=300)}

Usage Examples:
{self.fake.text(max_nb_chars=250)}

Troubleshooting:
{self.fake.text(max_nb_chars=200)}
"""
    
    def _generate_configuration(self) -> str:
        config_date = self.get_random_document_date()
        return f"""# Configuration File
# Generated on {config_date.strftime('%Y-%m-%d')}

[database]
host = {self.fake.ipv4()}
port = {self.fake.random_int(min=3000, max=9999)}
username = {self.fake.user_name()}
password = {self.fake.password()}

[server]
port = {self.fake.random_int(min=8000, max=9999)}
debug = {random.choice(['true', 'false'])}
max_connections = {self.fake.random_int(min=10, max=1000)}

[logging]
level = {random.choice(['DEBUG', 'INFO', 'WARNING', 'ERROR'])}
file = /var/log/{self.fake.word()}.log
"""


class SVGGenerator:
    """Generate SVG content for images"""
    
    def __init__(self):
        self.fake = Faker('en_US')
    
    def generate_svg(self, width: int, height: int) -> str:
        """Generate a proper SVG with correct viewBox and detailed paths"""
        # Ensure minimum size
        width = max(width, 100)
        height = max(height, 100)
        
        # Random background color
        bg_r, bg_g, bg_b = random.randint(240, 255), random.randint(240, 255), random.randint(240, 255)
        
        svg_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg width="{width}" height="{height}" viewBox="0 0 {width} {height}" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <linearGradient id="grad1" x1="0%" y1="0%" x2="100%" y2="100%">
      <stop offset="0%" style="stop-color:rgb({bg_r},{bg_g},{bg_b});stop-opacity:1" />
      <stop offset="100%" style="stop-color:rgb({bg_r-20},{bg_g-20},{bg_b-20});stop-opacity:1" />
    </linearGradient>
  </defs>
  <rect width="100%" height="100%" fill="url(#grad1)"/>
'''
        
        # Add detailed geometric shapes
        shapes = self._add_geometric_shapes(svg_content, width, height)
        for shape in shapes:
            svg_content += shape + '\n'
        
        # Add complex paths
        paths = self._add_curved_paths(svg_content, width, height)
        for path in paths:
            svg_content += path + '\n'
        
        # Add text elements
        texts = self._add_text_elements(svg_content, width, height)
        for text in texts:
            svg_content += text + '\n'
        
        # Add decorative elements
        decorations = self._add_decorative_elements(svg_content, width, height)
        for decoration in decorations:
            svg_content += decoration + '\n'
        
        svg_content += '</svg>'
        return svg_content
    
    def _add_geometric_shapes(self, svg_content: str, width: int, height: int):
        """Add various geometric shapes"""
        shapes = []
        
        # Add rectangles with rounded corners
        for _ in range(random.randint(2, 5)):
            x = random.randint(0, width - 100)
            y = random.randint(0, height - 100)
            w = random.randint(50, min(200, width - x))
            h = random.randint(30, min(150, height - y))
            rx = random.randint(0, 15)
            color = f"rgb({random.randint(100, 200)},{random.randint(100, 200)},{random.randint(100, 200)})"
            opacity = random.uniform(0.4, 0.8)
            
            shapes.append(f'  <rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" ry="{rx}" fill="{color}" opacity="{opacity}" stroke="rgb(80,80,80)" stroke-width="1"/>')
        
        # Add circles
        for _ in range(random.randint(1, 4)):
            cx = random.randint(50, width - 50)
            cy = random.randint(50, height - 50)
            r = random.randint(20, min(80, min(cx, width-cx, cy, height-cy)))
            color = f"rgb({random.randint(50, 180)},{random.randint(50, 180)},{random.randint(50, 180)})"
            opacity = random.uniform(0.3, 0.7)
            
            shapes.append(f'  <circle cx="{cx}" cy="{cy}" r="{r}" fill="{color}" opacity="{opacity}" stroke="rgb(60,60,60)" stroke-width="2"/>')
        
        # Add ellipses
        for _ in range(random.randint(1, 3)):
            cx = random.randint(60, width - 60)
            cy = random.randint(40, height - 40)
            rx = random.randint(30, min(100, min(cx, width-cx)))
            ry = random.randint(20, min(80, min(cy, height-cy)))
            color = f"rgb({random.randint(80, 220)},{random.randint(80, 220)},{random.randint(80, 220)})"
            opacity = random.uniform(0.2, 0.6)
            
            shapes.append(f'  <ellipse cx="{cx}" cy="{cy}" rx="{rx}" ry="{ry}" fill="{color}" opacity="{opacity}"/>')
        
        return shapes
    
    def _add_curved_paths(self, svg_content: str, width: int, height: int):
        """Add curved paths using SVG path elements"""
        paths = []
        
        for _ in range(random.randint(2, 4)):
            # Create a curved path
            start_x = random.randint(0, width // 3)
            start_y = random.randint(0, height // 3)
            
            # Control points for bezier curve
            cp1_x = random.randint(width // 4, 3 * width // 4)
            cp1_y = random.randint(0, height)
            cp2_x = random.randint(width // 4, 3 * width // 4)
            cp2_y = random.randint(0, height)
            
            end_x = random.randint(2 * width // 3, width)
            end_y = random.randint(2 * height // 3, height)
            
            path_data = f"M {start_x} {start_y} C {cp1_x} {cp1_y}, {cp2_x} {cp2_y}, {end_x} {end_y}"
            
            color = f"rgb({random.randint(40, 160)},{random.randint(40, 160)},{random.randint(40, 160)})"
            stroke_width = random.randint(2, 6)
            opacity = random.uniform(0.4, 0.8)
            
            paths.append(f'  <path d="{path_data}" stroke="{color}" stroke-width="{stroke_width}" fill="none" opacity="{opacity}"/>')
        
        return paths
    
    def _add_text_elements(self, svg_content: str, width: int, height: int):
        """Add text elements with proper positioning"""
        texts = []
        
        for _ in range(random.randint(1, 3)):
            x = random.randint(20, max(20, width - 150))
            y = random.randint(30, height - 20)
            font_size = random.randint(14, min(32, height // 10))
            
            text_options = [
                self.fake.word().title(),
                f"IMG_{random.randint(1000, 9999)}",
                self.fake.date_between(start_date='-10y', end_date='today').strftime('%Y-%m-%d'),
                self.fake.company()[:12],
                f"Document #{random.randint(100, 999)}",
                "SAMPLE",
                "DRAFT"
            ]
            
            text = random.choice(text_options)
            color = f"rgb({random.randint(20, 100)},{random.randint(20, 100)},{random.randint(20, 100)})"
            
            texts.append(f'  <text x="{x}" y="{y}" font-family="Arial, sans-serif" font-size="{font_size}" fill="{color}" font-weight="bold">{text}</text>')
        
        return texts
    
    def _add_decorative_elements(self, svg_content: str, width: int, height: int):
        """Add decorative elements like patterns and small shapes"""
        decorations = []
        
        # Add some small decorative circles
        for _ in range(random.randint(3, 8)):
            cx = random.randint(10, width - 10)
            cy = random.randint(10, height - 10)
            r = random.randint(3, 8)
            color = f"rgb({random.randint(150, 255)},{random.randint(150, 255)},{random.randint(150, 255)})"
            
            decorations.append(f'  <circle cx="{cx}" cy="{cy}" r="{r}" fill="{color}" opacity="0.6"/>')
        
        # Add some lines for texture
        for _ in range(random.randint(2, 5)):
            x1 = random.randint(0, width)
            y1 = random.randint(0, height)
            x2 = random.randint(0, width)
            y2 = random.randint(0, height)
            color = f"rgb({random.randint(180, 220)},{random.randint(180, 220)},{random.randint(180, 220)})"
            
            decorations.append(f'  <line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{color}" stroke-width="1" opacity="0.3"/>')
        
        return decorations


class SVGToImageConverter:
    """Convert SVG to various image formats using multiple methods"""
    
    @staticmethod
    def svg_to_png(svg_content: str, width: int, height: int) -> bytes:
        """Convert SVG to PNG using best available method"""
        
        # Method 1: Try cairosvg (most reliable)
        try:
            import cairosvg
            return cairosvg.svg2png(
                bytestring=svg_content.encode('utf-8'),
                output_width=width,
                output_height=height
            )
        except ImportError:
            pass
        except Exception as e:
            print(f"cairosvg failed: {e}")
        
        # Method 2: Try svglib + reportlab
        try:
            from svglib.svglib import renderSVG
            from reportlab.graphics import renderPM
            import tempfile
            
            # Write SVG to temp file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.svg', delete=False) as f:
                f.write(svg_content)
                svg_path = f.name
            
            try:
                # Convert SVG to ReportLab drawing
                drawing = renderSVG.renderSVG_to_file(svg_path)
                
                # Render to PNG
                buffer = io.BytesIO()
                renderPM.drawToFile(drawing, buffer, fmt='PNG')
                png_data = buffer.getvalue()
                
                os.unlink(svg_path)  # Clean up temp file
                return png_data
                
            except Exception:
                os.unlink(svg_path)  # Clean up temp file
                raise
                
        except ImportError:
            pass
        except Exception as e:
            print(f"svglib failed: {e}")
        
        # Method 3: Try wand (ImageMagick)
        try:
            from wand.image import Image as WandImage
            
            with WandImage(blob=svg_content.encode('utf-8'), format='svg') as img:
                img.format = 'png'
                img.resize(width, height)
                return img.blob
                
        except ImportError:
            pass
        except Exception as e:
            print(f"wand failed: {e}")
        
        # Method 4: Fallback to PIL with manual rendering
        return SVGToImageConverter._fallback_svg_to_png(svg_content, width, height)
    
    @staticmethod
    def _fallback_svg_to_png(svg_content: str, width: int, height: int) -> bytes:
        """Fallback method using PIL to create a simple image"""
        # Create a gradient background
        img = Image.new('RGB', (width, height), (240, 240, 240))
        draw = ImageDraw.Draw(img)
        
        # Create gradient effect
        for y in range(height):
            color_val = int(240 - (y / height) * 40)
            draw.line([(0, y), (width, y)], fill=(color_val, color_val, color_val + 10))
        
        # Add some geometric shapes based on common SVG elements
        for _ in range(random.randint(5, 12)):
            shape_type = random.choice(['rectangle', 'ellipse', 'line', 'polygon'])
            color = (random.randint(50, 200), random.randint(50, 200), random.randint(50, 200))
            
            if shape_type == 'rectangle':
                x1 = random.randint(0, width - 50)
                y1 = random.randint(0, height - 50)
                x2 = x1 + random.randint(20, min(200, width - x1))
                y2 = y1 + random.randint(20, min(150, height - y1))
                draw.rectangle([x1, y1, x2, y2], fill=color, outline=(color[0]//2, color[1]//2, color[2]//2))
                
            elif shape_type == 'ellipse':
                x1 = random.randint(20, width - 80)
                y1 = random.randint(20, height - 80)
                x2 = x1 + random.randint(40, min(160, width - x1))
                y2 = y1 + random.randint(30, min(120, height - y1))
                draw.ellipse([x1, y1, x2, y2], fill=color, outline=(color[0]//2, color[1]//2, color[2]//2))
                
            elif shape_type == 'line':
                x1 = random.randint(0, width)
                y1 = random.randint(0, height)
                x2 = random.randint(0, width)
                y2 = random.randint(0, height)
                draw.line([x1, y1, x2, y2], fill=color, width=random.randint(2, 5))
                
            elif shape_type == 'polygon':
                # Create triangle
                points = []
                center_x = random.randint(50, width - 50)
                center_y = random.randint(50, height - 50)
                radius = random.randint(20, 60)
                
                for i in range(3):
                    angle = (i * 120) * 3.14159 / 180
                    x = center_x + int(radius * math.cos(angle))
                    y = center_y + int(radius * math.sin(angle))
                    points.extend([x, y])
                
                draw.polygon(points, fill=color, outline=(color[0]//2, color[1]//2, color[2]//2))
        
        # Add text
        try:
            font_size = max(12, min(32, height // 20))
            fake = Faker('en_US')
            random_date = fake.date_between(start_date='-10y', end_date='today')
            text = f"Generated {random_date.strftime('%Y-%m-%d')}"
            draw.text((20, 20), text, fill=(50, 50, 50))
        except:
            pass
        
        # Convert to PNG bytes
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        return buffer.getvalue()


class JPGTemplate(BaseTemplate):
    """Generate realistic JPEG image files using SVG conversion"""
    
    def __init__(self, language: str = 'en'):
        super().__init__(language)
        self.svg_gen = SVGGenerator()
        self.converter = SVGToImageConverter()
    
    def generate(self, target_size: int, filename: str) -> bytes:
        # Generate appropriate dimensions
        width = random.randint(400, min(1920, max(800, target_size // 1000)))
        height = random.randint(300, min(1080, max(600, target_size // 1500)))
        
        # Generate SVG content
        svg_content = self.svg_gen.generate_svg(width, height)
        
        try:
            # Convert SVG to PNG first
            png_bytes = self.converter.svg_to_png(svg_content, width, height)
            
            # Convert PNG to JPEG
            img = Image.open(io.BytesIO(png_bytes))
            if img.mode in ('RGBA', 'LA', 'P'):
                # Convert to RGB for JPEG
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = rgb_img
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Save as JPEG with good quality
            buffer = io.BytesIO()
            quality = random.randint(75, 95)
            img.save(buffer, format='JPEG', quality=quality, optimize=True)
            content_bytes = buffer.getvalue()
            
        except Exception as e:
            print(f"SVG conversion failed: {e}")
            # Fallback: Create image directly with PIL
            content_bytes = self._create_fallback_jpeg(width, height)
        
        # Don't pad JPEG files as it corrupts them - return as is
        return content_bytes
    
    def _create_fallback_jpeg(self, width: int, height: int) -> bytes:
        """Create a fallback JPEG using PIL directly"""
        img = Image.new('RGB', (width, height), (random.randint(220, 255), random.randint(220, 255), random.randint(220, 255)))
        draw = ImageDraw.Draw(img)
        
        # Add gradient background
        for y in range(height):
            color_intensity = int(255 - (y / height) * 50)
            color = (color_intensity, color_intensity, min(255, color_intensity + 20))
            draw.line([(0, y), (width, y)], fill=color)
        
        # Add shapes
        for _ in range(random.randint(8, 15)):
            shape_type = random.choice(['rectangle', 'ellipse', 'line'])
            color = (random.randint(0, 180), random.randint(0, 180), random.randint(0, 180))
            
            if shape_type == 'rectangle':
                x1, y1 = random.randint(0, width-50), random.randint(0, height-50)
                x2, y2 = x1 + random.randint(30, 150), y1 + random.randint(20, 100)
                x2, y2 = min(x2, width), min(y2, height)
                draw.rectangle([x1, y1, x2, y2], fill=color, outline=(color[0]//2, color[1]//2, color[2]//2))
                
            elif shape_type == 'ellipse':
                x1, y1 = random.randint(0, width-60), random.randint(0, height-60)
                x2, y2 = x1 + random.randint(40, 120), y1 + random.randint(30, 100)
                x2, y2 = min(x2, width), min(y2, height)
                draw.ellipse([x1, y1, x2, y2], fill=color, outline=(color[0]//2, color[1]//2, color[2]//2))
                
            else:  # line
                x1, y1 = random.randint(0, width), random.randint(0, height)
                x2, y2 = random.randint(0, width), random.randint(0, height)
                draw.line([x1, y1, x2, y2], fill=color, width=random.randint(2, 6))
        
        # Add text
        try:
            fake = Faker('en_US')
            random_date = fake.date_between(start_date='-10y', end_date='today')
            text = f"IMG_{random.randint(1000, 9999)} - {random_date.strftime('%Y-%m-%d')}"
            draw.text((20, 20), text, fill=(50, 50, 50))
        except:
            pass
        
        buffer = io.BytesIO()
        img.save(buffer, format='JPEG', quality=85, optimize=True)
        return buffer.getvalue()


class PNGTemplate(BaseTemplate):
    """Generate realistic PNG image files using SVG conversion"""
    
    def __init__(self, language: str = 'en'):
        super().__init__(language)
        self.svg_gen = SVGGenerator()
        self.converter = SVGToImageConverter()
    
    def generate(self, target_size: int, filename: str) -> bytes:
        # Generate appropriate dimensions
        width = random.randint(400, min(1920, max(800, target_size // 800)))
        height = random.randint(300, min(1080, max(600, target_size // 1200)))
        
        # Generate SVG content
        svg_content = self.svg_gen.generate_svg(width, height)
        
        try:
            # Convert SVG to PNG directly
            content_bytes = self.converter.svg_to_png(svg_content, width, height)
            
        except Exception as e:
            print(f"SVG to PNG conversion failed: {e}")
            # Fallback: Create PNG directly with PIL
            content_bytes = self._create_fallback_png(width, height)
        
        # Don't pad PNG files as it corrupts them - return as is
        return content_bytes
    
    def _create_fallback_png(self, width: int, height: int) -> bytes:
        """Create a fallback PNG using PIL directly"""
        # Create image with transparency support
        img = Image.new('RGBA', (width, height), (random.randint(240, 255), random.randint(240, 255), random.randint(240, 255), 255))
        draw = ImageDraw.Draw(img)
        
        # Add gradient background
        for y in range(height):
            alpha = int(255 - (y / height) * 50)
            color_intensity = int(250 - (y / height) * 30)
            color = (color_intensity, color_intensity, min(255, color_intensity + 15), alpha)
            draw.line([(0, y), (width, y)], fill=color)
        
        # Add shapes with transparency
        for _ in range(random.randint(6, 12)):
            shape_type = random.choice(['rectangle', 'ellipse', 'line', 'polygon'])
            alpha = random.randint(120, 200)
            color = (random.randint(0, 180), random.randint(0, 180), random.randint(0, 180), alpha)
            
            if shape_type == 'rectangle':
                x1, y1 = random.randint(0, width-60), random.randint(0, height-60)
                x2, y2 = x1 + random.randint(40, 160), y1 + random.randint(30, 120)
                x2, y2 = min(x2, width), min(y2, height)
                draw.rectangle([x1, y1, x2, y2], fill=color, outline=(color[0]//2, color[1]//2, color[2]//2, 255))
                
            elif shape_type == 'ellipse':
                x1, y1 = random.randint(0, width-80), random.randint(0, height-80)
                x2, y2 = x1 + random.randint(50, 140), y1 + random.randint(40, 120)
                x2, y2 = min(x2, width), min(y2, height)
                draw.ellipse([x1, y1, x2, y2], fill=color, outline=(color[0]//2, color[1]//2, color[2]//2, 255))
                
            elif shape_type == 'line':
                x1, y1 = random.randint(0, width), random.randint(0, height)
                x2, y2 = random.randint(0, width), random.randint(0, height)
                draw.line([x1, y1, x2, y2], fill=color, width=random.randint(3, 8))
                
            elif shape_type == 'polygon':
                # Create triangle with transparency
                points = []
                center_x = random.randint(60, width - 60)
                center_y = random.randint(60, height - 60)
                radius = random.randint(30, 80)
                
                for i in range(3):
                    angle = (i * 120) * math.pi / 180
                    x = center_x + int(radius * math.cos(angle))
                    y = center_y + int(radius * math.sin(angle))
                    points.extend([x, y])
                
                draw.polygon(points, fill=color, outline=(color[0]//2, color[1]//2, color[2]//2, 255))
        
        # Add text with transparency
        try:
            fake = Faker('en_US')
            random_date = fake.date_between(start_date='-10y', end_date='today')
            text = f"PNG_{random.randint(1000, 9999)} - {random_date.strftime('%Y-%m-%d')}"
            draw.text((20, 20), text, fill=(30, 30, 30, 200))
        except:
            pass
        
        buffer = io.BytesIO()
        img.save(buffer, format='PNG', optimize=True)
        return buffer.getvalue()


class TemplateFactory:
    """Factory class for creating template instances"""
    
    template_classes = {
        'pdf': PDFTemplate,
        'docx': DOCXTemplate,
        'xlsx': XLSXTemplate,
        'eml': EMLTemplate,
        'csv': CSVTemplate,
        'txt': TXTTemplate,
        'jpg': JPGTemplate,
        'png': PNGTemplate,
    }
    
    @classmethod
    def create_template(cls, file_type: str, language: str = 'en') -> BaseTemplate:
        """Create a template instance for the given file type"""
        template_class = cls.template_classes.get(file_type.lower())
        if not template_class:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return template_class(language=language)
    
    @classmethod
    def get_supported_types(cls) -> List[str]:
        """Get list of supported file types"""
        return list(cls.template_classes.keys())